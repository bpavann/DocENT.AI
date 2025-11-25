from typing import List
from docx import Document
from pdfdocument.document import PDFDocument
import os

class ExportTools:
    def to_docx(self, content, filename="DocENT.AI.docx") -> str:
        """Export content to DOCX with fixed user-visible filename."""
        doc = Document()
        doc.add_paragraph(str(content))
        doc.save(filename)
        return filename  # Return user-facing file name

    def to_pdf(self, content, filename="DocENT.AI.pdf") -> str:
        """Export content to PDF with fixed user-visible filename."""
        pdf = PDFDocument(filename)
        pdf.init_report()

        pdf.h1("DocENT AI Generated Report")
        pdf.p("This is a PDF generated using DocENT.AI.")

        for line in str(content).split("\n"):
            pdf.p(line)

        pdf.generate()
        return filename  # Return user-facing file name


class AutomationAgent:
    def __init__(self):
        self.export_tools = ExportTools()

    def export(self, content, formats: List[str] = None) -> dict:
        """Export content to DOCX and/or PDF using fixed file name DocENT.AI."""
        if formats is None:
            formats = ["docx", "pdf"]

        results = {}
        for fmt in formats:
            fmt_lower = fmt.lower()
            if fmt_lower == "docx":
                results["docx"] = self.export_tools.to_docx(content)
            elif fmt_lower == "pdf":
                results["pdf"] = self.export_tools.to_pdf(content)
            else:
                results[fmt_lower] = None  # unsupported format
        return results


# Example usage
if __name__ == "__main__":
    agent = AutomationAgent()
    content = "This is a sample report.\nGenerated for testing.\nSupports multiple lines."
    exported_files = agent.export(content)
    print("Files available for user download:", exported_files)
